from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_lang(document: Document, lang: str = 'es-UY') -> None:
    styles = document.styles
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Title']:
        try:
            style = styles[style_name]
        except KeyError:
            continue
        rpr = style.element.get_or_add_rPr()
        lang_el = rpr.find(qn('w:lang'))
        if lang_el is None:
            lang_el = OxmlElement('w:lang')
            rpr.append(lang_el)
        lang_el.set(qn('w:val'), lang)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def process_inline(paragraph, text: str) -> None:
    # basic support for inline code, bold, italics, and markdown links
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\)|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('['):
            m2 = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', token)
            label = m2.group(1) if m2 else token
            url = m2.group(2) if m2 else ''
            paragraph.add_run(label)
            if url:
                run = paragraph.add_run(f' ({url})')
                run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines: list[str], start: int):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].rstrip())
        i += 1
    return table_lines, i


def split_md_row(line: str) -> list[str]:
    line = line.strip().strip('|')
    return [cell.strip() for cell in line.split('|')]


def add_md_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return
    headers = split_md_row(table_lines[0])
    rows = [split_md_row(line) for line in table_lines[2:]]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for idx, cell_text in enumerate(headers):
        cell_p = table.rows[0].cells[idx].paragraphs[0]
        process_inline(cell_p, cell_text)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            cell_p = table.rows[r_idx].cells[c_idx].paragraphs[0]
            process_inline(cell_p, cell_text)
    doc.add_paragraph()


def add_image(doc: Document, base_dir: Path, alt: str, raw_path: str) -> None:
    path = Path(raw_path)
    if not path.is_absolute():
        path = (base_dir / path).resolve()
    if not path.exists():
        p = doc.add_paragraph()
        process_inline(p, f'[Imagen no encontrada: {alt} -> {path}]')
        return
    if path.suffix.lower() == '.svg':
        p = doc.add_paragraph()
        process_inline(p, f'[Figura no incrustada por limitación de formato SVG en esta conversión: {alt}. Archivo: {path}]')
        return
    try:
        doc.add_picture(str(path), width=Inches(6.0))
        cap = doc.add_paragraph()
        process_inline(cap, alt)
        cap.style = 'Caption'
    except Exception as exc:
        p = doc.add_paragraph()
        process_inline(p, f'[Error al incrustar imagen: {alt} -> {path} ({exc})]')


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    set_lang(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)

    i = 0
    in_code = False
    code_buffer: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                p = doc.add_paragraph(style='No Spacing')
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Courier New'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(9)
                in_code = False
                code_buffer = []
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped == '---':
            add_page_break(doc)
            i += 1
            continue

        img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', stripped)
        if img_match:
            add_image(doc, md_path.parent, img_match.group(1), img_match.group(2))
            i += 1
            continue

        if stripped.startswith('|'):
            table_lines, i = parse_table(lines, i)
            add_md_table(doc, table_lines)
            continue

        heading = re.match(r'^(#{1,6})\s+(.*)$', line)
        if heading:
            level = min(len(heading.group(1)), 4)
            text = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style='Title')
            else:
                p = doc.add_paragraph(style=f'Heading {level}')
            process_inline(p, text)
            i += 1
            continue

        if stripped.startswith('>'):
            p = doc.add_paragraph(style='Intense Quote')
            process_inline(p, stripped.lstrip('> ').strip())
            i += 1
            continue

        bullet = re.match(r'^[-*]\s+(.*)$', stripped)
        if bullet:
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, bullet.group(1))
            i += 1
            continue

        ordered = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if ordered:
            p = doc.add_paragraph(style='List Number')
            process_inline(p, ordered.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        process_inline(p, line)
        i += 1

    doc.save(str(docx_path))


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('usage: markdown_to_docx.py INPUT.md OUTPUT.docx')
        raise SystemExit(2)
    convert(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
